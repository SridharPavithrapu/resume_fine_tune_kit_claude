# generate_with_models.py
from __future__ import annotations
import os, re, json, hashlib, time
from typing import Dict, List, Optional, Tuple, Any
from pathlib import Path as _Path
from docx import Document
from resume_tailoring.utils import load_resume_text

def _hard_replace_section(full_text: str, section_name: str, new_body: str) -> str:
    """
    Replace the entire named section with `new_body` while keeping the section header line.
    Boundary = next ALL-CAPS header or EOF.
    """
    header_pat = rf"(?m)^{re.escape(section_name)}\s*\n"
    m = re.search(header_pat, full_text)
    if not m:
        return f"{section_name}\n\n{new_body.rstrip()}\n\n" + full_text
    next_m = re.search(r"(?m)^\s*[A-Z][A-Z &/]{2,}\s*$", full_text[m.end():])
    end = m.end() + next_m.start() if next_m else len(full_text)
    return full_text[:m.end()] + new_body.rstrip() + "\n" + full_text[end:]


def _keep_only_bullets_and_blank(text: str) -> str:
    """Remove any non-bullet, non-blank lines (e.g., headers that slipped in)."""
    lines = []
    for ln in (text or "").splitlines():
        s = ln.rstrip()
        if not s or s.startswith("- "):
            lines.append(s)
        # else: drop header/stray lines
    # preserve existing blank lines
    return "\n".join(lines)

def _anchors_from_block(block: str) -> set[str]:
    import re
    tokens = re.findall(r"[A-Za-z][A-Za-z0-9+\-#\.]{2,}", block or "")
    hot = {"tableau","qlik","snowflake","sql","ssis","sdlc","etl","bi","dashboard",
           "compliance","audit","risk","claims","provider","hipaa","warehouse","sap",
           "procurement","logistics","supply","dimensional","modeling","python","excel"}
    anchors = set()
    for t in tokens:
        lt = t.lower()
        if lt in hot or any(ch.isupper() for ch in t):
            anchors.add(lt)
    return anchors

def _score_by_overlap(lines: list[str], want: set[str]) -> list[tuple[int,str]]:
    scored = []
    want = {w.lower() for w in (want or set())}
    for ln in (lines or []):
        l = ln.lower()
        score = sum(1 for a in want if a in l)
        scored.append((score, ln))
    return sorted(scored, key=lambda x: (-x[0], len(x[1])))

def _dedup_fuzzy(lines: list[str], thresh: float = 0.75) -> list[str]:
    import re
    out = []
    def tok(s): 
        return set(re.findall(r"[a-z0-9]+", (s or "").lower()))
    for ln in (lines or []):
        t = tok(ln)
        if not t:
            continue
        keep = True
        for have in out:
            th = tok(have)
            j = len(t & th) / max(1, len(t | th))
            if j >= thresh:
                keep = False; break
        if keep:
            out.append(ln)
    return out

# ============== Debug helpers (pre-flight + general) ==============
def _dbg_write(name: str, content: str):
    d = _Path("prompt_logs"); d.mkdir(exist_ok=True)
    with (d / name).open("w", encoding="utf-8") as f:
        f.write(content if content is not None else "")

def _redact_pii(text: str) -> str:
    if not text: return text
    text = re.sub(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", "[redacted@email]", text)
    text = re.sub(r"\+?\d(?:[\s\-().]*\d){7,}", "[redacted phone]", text)
    return text

def _quick_stats(label: str, s: str) -> dict:
    s = s or ""
    w = len(re.findall(r"\S+", s))
    lines = s.count("\n") + (1 if s else 0)
    approx_tokens = max(1, len(s)//4)  # coarse estimate
    md5 = hashlib.md5(s.encode("utf-8")).hexdigest()[:12]
    return {"label": label, "chars": len(s), "words": w, "lines": lines, "approx_tokens": approx_tokens, "md5": md5}

# ---------------- keyword sanitizer (belt & suspenders) ----------------
SAFE_CHARS = set("abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789 +#./&()-_")
BLOCK_TOKENS = {"@media","var(","calc(","rgba(","webkit","moz","svg","inline","fa-","{","}","</",">","script","style"}

def _is_reasonable_keyword(kw: str) -> bool:
    if not kw: return False
    s = kw.strip()
    if len(s) < 2 or len(s) > 40: return False
    low = s.lower()
    if any(t in low for t in BLOCK_TOKENS): return False
    if not any(ch.isalnum() for ch in s): return False
    if any(c not in SAFE_CHARS for c in s): return False
    return True

def _sanitize_keywords(keywords: List[str]):
    keep, drop, seen = [], [], set()
    for k in (keywords or []):
        k = (k or "").strip()
        if not k: continue
        lk = k.lower()
        if lk in seen: continue
        if _is_reasonable_keyword(k):
            keep.append(k); seen.add(lk)
        else:
            drop.append(k)
    return keep[:80], drop

# ------------- Lightweight fallbacks to avoid hard deps -------------
try:
    from utils import (
        extract_resume_sections,
        insert_or_replace_section,
        parse_bullet_output,
    )
except Exception:
    def extract_resume_sections(text: str) -> Dict[str, str]:
        sections = {}
        current = None
        buf: List[str] = []
        for line in text.splitlines():
            s = line.strip().upper()
            if s in {"SUMMARY", "SKILLS", "WORK EXPERIENCE", "CERTIFICATIONS", "EDUCATION"}:
                if current is not None:
                    sections[current] = "\n".join(buf).strip("\n")
                current = s
                buf = []
            else:
                buf.append(line)
        if current is not None:
            sections[current] = "\n".join(buf).strip("\n")
        return sections

    def insert_or_replace_section(text: str, header: str, content: str) -> str:
        lines = text.splitlines()
        H = header.strip().upper()
        out: List[str] = []
        i, n, found = 0, len(lines), False
        while i < n:
            if lines[i].strip().upper() == H:
                found = True
                out += [header, ""]
                if content:
                    out.append(content.rstrip())
                i += 1
                while i < n and lines[i].strip().upper() not in {"SUMMARY","SKILLS","WORK EXPERIENCE","CERTIFICATIONS","EDUCATION"}:
                    i += 1
                continue
            out.append(lines[i]); i += 1
        if not found:
            if out and out[-1].strip():
                out.append("")
            out += [header, ""]
            if content:
                out.append(content.rstrip())
        return "\n".join(out)

    # permissive fallback parser: accepts "- " bullets OR plain sentence lines
    def parse_bullet_output(block: str) -> List[str]:
        lines = [ln.rstrip() for ln in block.splitlines()]
        bullets = [ln[2:].strip() for ln in lines if ln.strip().startswith("- ")]
        if bullets:
            return bullets
        out = []
        for ln in lines:
            s = ln.strip()
            if not s:
                continue
            if re.search(r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)\s+\d{4}", s, re.I):
                continue
            if re.match(r"^[A-Za-z].{20,}", s):  # sentence-like
                out.append(s)
        return out

# ---------------- Guard/cleanup helpers ----------------
try:
    from guard_clean_resume import (
        fix_broken_headers,
        remove_hallucinated_titles,
        remove_placeholder_bullets,
    )
except Exception:
    def fix_broken_headers(t: str) -> str: return t
    def remove_hallucinated_titles(t: str) -> str: return t
    def remove_placeholder_bullets(t: str) -> str: return t

DEBUG = os.environ.get("DEBUG_RESUME", "0") == "1"

# ---------------- Header detection & safety ----------------
JOB_HEADER_FULL = re.compile(r"""^
(?P<title>[A-Za-z][A-Za-z0-9/&().,+\- ]{1,80}?)\s*[–—-]\s*
(?P<company>[^,\n]{2,})(?:,\s*[^\n]*?)?
.*?\b((Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\s+\d{4}|(19|20)\d{2}|Present)\b
.*$""", re.IGNORECASE | re.VERBOSE)

JOB_HEADER_TRUNC = re.compile(r"^(?P<title>[A-Za-z0-9/&().,+\- ]*?)\s*[–—-]\s*$")

def is_job_header(line: str) -> bool:
    s = line.strip()
    m = JOB_HEADER_FULL.match(s)
    if not m:
        return False
    title = m.group("title").strip()
    if len(title.split()) > 8:
        return False
    return True

def _collect_base_headers(base_text: str) -> List[str]:
    return [ln.rstrip() for ln in base_text.splitlines() if is_job_header(ln)]

def _norm_title(s: str) -> str:
    return re.sub(r"\s+", " ", s.strip().lower())

def restore_missing_company_headers(text: str, expected_headers: List[str]) -> str:
    title_to_header: Dict[str, str] = {}
    for h in expected_headers:
        m = JOB_HEADER_FULL.match(h.strip())
        if not m: continue
        key = _norm_title(m.group("title"))
        if key not in title_to_header or len(h) > len(title_to_header[key]):
            title_to_header[key] = h
    out: List[str] = []
    for ln in text.splitlines():
        mt = JOB_HEADER_TRUNC.match(ln.strip())
        if mt:
            full = title_to_header.get(_norm_title(mt.group("title")))
            if full:
                out.append(full); continue
        out.append(ln)
    return "\n".join(out)

# ---------------- Model output cleaning ----------------
def _strip_md_separators(text: str) -> str:
    return re.sub(r"(?m)^\s*(?:---+|\*\*\*+|___+)\s*$", "", text)

def _clean_model_block(raw: str) -> str:
    cleaned = raw or ""
    cleaned = re.sub(r"(?m)^\s*```.*$", "", cleaned)
    cleaned = re.sub(r"(?m)^\s*#{1,6}\s+.*$", "", cleaned)
    cleaned = re.sub(r"(?m)^\s*\*{2}.+?\*{2}\s*$", "", cleaned)
    cleaned = _strip_md_separators(cleaned)
    cleaned = re.sub(r"(?im)^\s*(rewritten|final|output)\s*:\s*", "", cleaned)
    cleaned = re.sub(r"(?m)^\s*[•●▪︎➤➔→►]\s*", "- ", cleaned)
    cleaned = re.sub(r"(?m)^\s*[\*-]\s+", "- ", cleaned)
    cleaned = re.sub(r"(?m)^\s*[–—]\s+", "- ", cleaned)
    cleaned = re.sub(r"(?m)^\s*\d+\.\s+", "- ", cleaned)
    lines = cleaned.splitlines()
    first_bullet_idx = next((i for i,l in enumerate(lines) if l.strip().startswith("- ")), None)
    if first_bullet_idx is not None:
        cleaned = "\n".join(lines[first_bullet_idx:])
    cleaned = re.sub(r"\*\*\s*$", "", cleaned.strip())
    return cleaned.strip()

def _clean_summary_output(text: str) -> str:
    if not text: return ""
    # Strip LLM fluff like "Here is a (tailored|possible) summary..." etc.
    text = re.sub(
        r"(?im)^\s*(here\s+is\s+(a|the)\s+(tailored|possible)\s+resume\s+summary.*?:|here\s+is\s+(a|the)\s+summary.*?:)\s*",
        "", text.strip()
    )
    text = re.sub(r"(?im)^\s*(summary\s*:)\s*", "", text)
    text = re.sub(r"(?m)^\s*#{1,6}\s+.*$", "", text)
    # Strip stray quotes
    return text.strip().strip('"\u201C\u201D').strip()


# ----- New instruction-line scrubber + job-aware validators -----
INSTR_LINE = re.compile(r"(?im)^\s*(?:-?\s*)?(use|follow|keep|write|stick|ensure|remember|avoid)\b.*")
SEPARATOR_LINE = re.compile(r"(?im)^\s*-{3,}\s*$")

def _clean_instruction_lines(text: str) -> str:
    if not text:
        return ""
    parts = text.splitlines()
    out = []
    seen_bullet = False
    for ln in parts:
        s = ln.rstrip()
        if not seen_bullet:
            if s.strip().startswith("- "):
                seen_bullet = True
            else:
                continue  # skip preface noise
        if INSTR_LINE.match(s) or SEPARATOR_LINE.match(s):
            continue
        out.append(s)
    cleaned = "\n".join(out).strip()
    cleaned = re.sub(r"(?im)^\s*(use|follow|keep|write|stick|ensure|remember|avoid)\b.*$", "", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned).strip()
    return cleaned

def _split_jobs_and_count_bullets(text: str) -> list[int]:
    if not text: return []
    blocks = [b for b in re.split(r"(?:\n\s*){2,}", text.strip()) if b.strip()]
    return [sum(1 for ln in blk.splitlines() if ln.strip().startswith("- ")) for blk in blocks]

def _jobwise_ok(text: str, expected_jobs: int, lo: int = 5, hi: int = 6) -> bool:
    counts = _split_jobs_and_count_bullets(text)
    if expected_jobs and len(counts) != expected_jobs:
        return False
    return bool(counts) and all(lo <= c <= hi for c in counts)

def _extract_bullets_per_job(model_text: str, max_bullets: int = 6) -> List[List[str]]:
    txt = _clean_model_block(model_text)
    txt = _strip_md_separators(txt)
    blocks = re.split(r"(?:\n\s*){2,}", txt)
    jobs: List[List[str]] = []
    for blk in blocks:
        blk = re.sub(r"(?m)^\s*#{1,6}\s+.*$", "", blk)
        blk = re.sub(r"(?m)^\s*\*{2}.+?\*{2}\s*$", "", blk)
        raw = [ln.strip() for ln in blk.splitlines() if ln.strip().startswith("- ")]
        seen, bullets = set(), []
        for ln in raw:
            b = ln[2:].strip().rstrip(".")
            key = b.lower()
            if key and key not in seen:
                seen.add(key)
                bullets.append(b + ".")
        if bullets:
            jobs.append(bullets[:max_bullets])
    return jobs

def _enforce_bullet_quality(bullets: List[str]) -> List[str]:
    out = []
    for b in bullets:
        s = re.sub(r"^[•\-\u2022]\s*", "", b.strip().rstrip("."))
        if s: out.append(s + ".")
    seen, dedup = set(), []
    for b in out:
        if b.lower() in seen: continue
        seen.add(b.lower()); dedup.append(b)
    return dedup

def _looks_sentence_like(line: str) -> bool:
    s = line.strip()
    if not s or is_job_header(s): return False
    if re.search(r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)\s+\d{4}", s, re.I):
        return False
    return bool(re.match(r"^[A-Za-z].{20,}", s))

# ---- Additional validator to catch junk outputs like "Use **bold** ..." ----
_INSTRUCTION_PAT = re.compile(r"(?i)\b(use|keep|focus|do not|don't|bold|separate|headers?)\b")
def _is_valid_bullet_block(txt: str, min_bullets: int = 4) -> bool:
    if not txt or not txt.strip():
        return False
    lines = [l.strip() for l in txt.splitlines() if l.strip()]
    bullets = [l for l in lines if l.startswith("- ")]
    if len(bullets) < min_bullets:
        return False
    for b in bullets:
        if _INSTRUCTION_PAT.search(b[2:]) and not re.search(r"[0-9%$]", b):
            return False
        if not re.search(r"[A-Za-z]", b):
            return False
    return True

# ---- Blank line hygiene between sections ----
def _ensure_blank_lines_between_sections(text: str) -> str:
    VALID = {"SUMMARY","SKILLS","WORK EXPERIENCE","CERTIFICATIONS","EDUCATION"}
    lines = text.splitlines()
    out: List[str] = []
    for i, ln in enumerate(lines):
        if ln.strip().upper() in VALID:
            if out and out[-1].strip():
                out.append("")  # ensure gap before each header
        out.append(ln)
    return "\n".join(out)

# ---- Certifications cleanup (drop placeholder dashes only) ----
def _strip_placeholder_block(s: str) -> str:
    if not s: return ""
    lines = [ln.strip() for ln in s.splitlines() if ln.strip()]
    if not lines: return ""
    if all(re.fullmatch(r"[–—\-]+", ln) for ln in lines):
        return ""
    return s.strip()

# ---- Rebalance bullets if model returns a single mega-block ----
def _rebalance_bullets(job_bullets: List[List[str]], num_jobs: int, min_per: int = 5, max_per: int = 6) -> List[List[str]]:
    if num_jobs <= 0:
        return job_bullets
    if len(job_bullets) == num_jobs:
        # already aligned
        return [lst[:max_per] for lst in job_bullets]
    if len(job_bullets) == 1 and num_jobs > 1:
        pool = job_bullets[0]
        out: List[List[str]] = []
        idx = 0
        for j in range(num_jobs):
            remaining_jobs = (num_jobs - j - 1)
            remaining_items = max(0, len(pool) - idx)
            need_min_for_rest = remaining_jobs * min_per
            take = min(max_per, max(min_per, remaining_items - need_min_for_rest))
            if remaining_items <= 0:
                out.append([])
                continue
            out.append(pool[idx: idx + take])
            idx += take
        return out
    # pad or trim to match job count
    flat = [b for blk in job_bullets for b in blk]
    out: List[List[str]] = []
    idx = 0
    for j in range(num_jobs):
        out.append(flat[idx: idx + max_per])
        idx += max_per
    return out

# ------------- Rebuild WORK EXPERIENCE safely -------------

def _chunk_bullets_forgiving(raw: str) -> list[list[str]]:
    """Forgiving splitter for headerless Together output.
    Splits on (a) blank lines, (b) dashed/underscore lines, (c) markdown separators.
    Returns groups of cleaned "- " bullets.
    """
    if not raw:
        return []
    txt = re.sub(r"(?m)^\s*(?:[-_]{3,}|\*{3,}|={3,})\s*$", "", raw.strip())
    parts = re.split(r"(?:\n\s*\n)+", txt)
    groups: list[list[str]] = []
    for blk in parts:
        lines = [ln.strip() for ln in blk.splitlines() if ln.strip().startswith("- ")]
        seen, bullets = set(), []
        for ln in lines:
            body = re.sub(r"^[•\-\u2022]\s*", "", ln).strip().rstrip(".")
            key = body.lower()
            if key and key not in seen:
                seen.add(key)
                bullets.append(body + ".")
        if bullets:
            groups.append(bullets)
    return groups

def fix_work_experience_layout(rewritten_block: str, base_text: str) -> str:
    """Map headerless bullets back to original job headers, with rebalance, anchors, and dedup."""
    base_headers = _collect_base_headers(base_text)
    if DEBUG:
        print(f"🧭 Base headers found: {len(base_headers)}")
        for i, h in enumerate(base_headers, 1):
            print(f"  {i}. {h}")

    # Strict parse first
    job_bullets = _extract_bullets_per_job(rewritten_block)

    # If strict parse weak (0/1 groups or few bullets), try forgiving chunking
    if not job_bullets or sum(len(g) for g in job_bullets) < 5 or len(job_bullets) == 1:
        forgiving = _chunk_bullets_forgiving(rewritten_block)
        if forgiving:
            job_bullets = forgiving

    if DEBUG:
        total = sum(len(g) for g in job_bullets) if job_bullets else 0
        print(f"🧪 Model job blocks detected (post-forgiving): {len(job_bullets)}; bullets total={total}")

    # Rebalance to match header count
    if len(job_bullets) != len(base_headers):
        job_bullets = _rebalance_bullets(job_bullets, len(base_headers), min_per=5, max_per=6)

    # Build base groups (for top-up)
    base_sections = extract_resume_sections(base_text)
    base_we = base_sections.get("WORK EXPERIENCE", base_text)
    groups: list[tuple[str, list[str]]] = []
    cur_header, buf = None, []
    for ln in base_we.splitlines():
        if is_job_header(ln):
            if cur_header is not None:
                blk = "\n".join(buf)
                cand = parse_bullet_output(blk) or [x.strip() for x in blk.splitlines() if _looks_sentence_like(x)]
                groups.append((cur_header, cand))
            cur_header = ln.rstrip()
            buf = []
        else:
            buf.append(ln)
    if cur_header is not None:
        blk = "\n".join(buf)
        cand = parse_bullet_output(blk) or [x.strip() for x in blk.splitlines() if _looks_sentence_like(x)]
        groups.append((cur_header, cand))

    out: list[str] = []
    jd_terms = set(globals().get("_GLOBAL_JD_KEYWORDS", set()))
    for idx, header in enumerate(base_headers):
        out += [header, ""]
        cand = _enforce_bullet_quality(job_bullets[idx]) if idx < len(job_bullets) else []

        # Top-up from base until 5; cap at 6
        base_cand = _enforce_bullet_quality(groups[idx][1]) if idx < len(groups) else []
        seen = {b.lower() for b in cand}
        for b in base_cand:
            if len(cand) >= 5:
                break
            if b.lower() not in seen:
                cand.append(b); seen.add(b.lower())

        # Enforce job-specific anchor coverage (>=2 bullets hit this job's anchors)
        job_block_text = "\n".join(base_cand)
        job_anchors = _anchors_from_block(job_block_text)
        want = (job_anchors or set()) | jd_terms
        coverage = sum(1 for b in cand if any(a in b.lower() for a in job_anchors))
        if coverage < 2 and base_cand:
            ranked = [ln for _, ln in _score_by_overlap(base_cand, want)]
            for ln in ranked:
                if len(cand) >= 6 or coverage >= 2:
                    break
                if ln.lower() not in {c.lower() for c in cand}:
                    cand.append(ln)
                    if any(a in ln.lower() for a in job_anchors):
                        coverage += 1

        # Final fuzzy de-dup + cap
        cand = _dedup_fuzzy(cand)[:6]

        for b in cand:
            out.append(f"- {b}")
        out.append("")

    return "\n".join(out).rstrip() + "\n"

def _call_summary_model(groq_client: Any, jd_text: str, keywords: List[str]) -> str:
    prompt = (
        "You are tailoring a resume SUMMARY for a specific job. "
        "Write 4–5 impactful lines (paragraph, not bullets) that integrate these keywords naturally (no list):\n"
        f"Keywords: {', '.join(keywords[:20])}\n\n"
        "Tone: crisp, metrics-driven, business impact. Avoid boilerplate and repetition."
    )
    return _clean_model_block(groq_client.complete(model=os.environ.get("GROQ_MODEL","llama-3.1-8b-instant"), prompt=prompt))


def _call_bullet_model(together_client: Any, base_we_text: str, jd_text: str, keywords: List[str]) -> str:
    """Plain-text bullets only; 5–6 per job; ≥3 quantified; strong verbs; no headers/dates."""
    sys = (
        "You rewrite WORK EXPERIENCE bullets job-by-job. Keep original job order and EXACT number of jobs.\n"
        'Respond with PLAIN TEXT ONLY: dash bullets ("- ...") with ONE blank line between jobs.\n'
        "Do NOT include headings, job titles, company names, dates, markdown, or any instructions/prefaces."
    )
    user = (
        "Rewrite the bullets below so they directly match the job description and the focus terms.\n"
        "Constraints:\n"
        "- Each job must have 5–6 bullets; every line begins with '- '\n"
        "- Exactly one blank line between job blocks\n"
        "- NO headers, NO titles, NO companies, NO dates, NO separators like '---'\n"
        "- At least 3 bullets per job include numbers (%, $, #, time) to quantify impact\n"
        "- Use strong action verbs and outcome-first phrasing (What → How → Impact)\n"
        "- Significantly rewrite content; avoid copying original lines verbatim\n"
        "- Naturally incorporate as many of the focus terms as relevant; avoid keyword stuffing\n"
        "----\n"
        "Job Description:\n" + jd_text.strip() + "\n"
        "----\n"
        "Focus terms (prioritize these where relevant):\n" + ", ".join(keywords[:50]) + "\n"
        "----\n"
        "Original WORK EXPERIENCE bullets (headers excluded):\n" + base_we_text.strip() + "\n"
        "----\n"
        "Return ONLY the bullets."
    )
    raw = together_client.chat(
        system=sys,
        user=user,
        model=os.environ.get("TOGETHER_MODEL","mistralai/Mixtral-8x7B-Instruct-v0.1"),
    )
    return raw or ""

# --------------------------- Orchestrator ---------------------------

def _extract_section(text: str, section_name: str) -> str:
    """Return the raw body of a named ALL-CAPS section (without the header line)."""
    m = re.search(rf"(?ms)^{re.escape(section_name)}\s*\n(.*?)(\n(?=[A-Z][A-Z &/]+\s*$)|\Z)", text)
    return m.group(1) if m else ""

def _norm_block(s: str) -> str:
    """Normalize a text block for robust comparison (whitespace + bullets)."""
    s = (s or "").strip()
    s = re.sub(r"\r\n", "\n", s)
    s = re.sub(r"\n\s*\n+", "\n\n", s)            # collapse multiple blank lines
    s = re.sub(r"^[•\-\u2022]\s*", "- ", s, flags=re.M)  # normalize bullet marks
    s = re.sub(r"\s+", " ", s)                    # collapse spaces
    return s.strip().lower()

def _tailor_resume_with_models_core(
    base_text: str,
    jd_text: str,
    all_keywords: List[str],
    *,
    groq_client: Any,
    together_client: Any,
    force_education: bool = True,
) -> str:
    sections = extract_resume_sections(base_text)

    # Sanitize keywords (observability: keep logs)
    all_keywords, dropped = _sanitize_keywords(all_keywords)
    _dbg_write("keywords_used.txt", "\n".join(all_keywords))
    global _GLOBAL_JD_KEYWORDS
    _GLOBAL_JD_KEYWORDS = {kw.lower() for kw in all_keywords}
    if dropped:
        _dbg_write("keywords_dropped.txt", "\n".join(dropped))

    # 1) SUMMARY
    try:
        summary_text = _call_summary_model(groq_client, jd_text, all_keywords).strip()
    except Exception:
        summary_text = sections.get("SUMMARY","")
    # Clean fluff if any
    summary_text = _clean_summary_output(summary_text)
    if not summary_text:
        summary_text = sections.get("SUMMARY","")

    # 2) WORK EXPERIENCE — PRE-FLIGHT LOGGING + NORMALIZATION
    base_we = sections.get("WORK EXPERIENCE","")

    # a) Detect headers to preserve
    base_headers = _collect_base_headers(base_text)
    _dbg_write("we_input_headers.txt", "\n".join(base_headers))

    # b) Normalize WE so sentence lines become "- " bullets for model input (does not modify final resume directly)
    we_lines = base_we.splitlines()
    norm_lines = []
    for ln in we_lines:
        if ln.strip().startswith("- "):
            norm_lines.append(ln)
        elif _looks_sentence_like(ln):
            norm_lines.append(f"- {ln.strip()}")
        else:
            norm_lines.append(ln)
    base_we_normalized = "\n".join(norm_lines)
    base_we_for_model = _keep_only_bullets_and_blank(base_we_normalized)  # define BEFORE preview


    # c) Save raw & normalized WE, JD, keywords, and prompt preview
    _dbg_write("we_input_raw.txt", _redact_pii(base_we))
    _dbg_write("we_input_normalized.txt", _redact_pii(base_we_normalized))
    _dbg_write("jd_input.txt", _redact_pii(jd_text))

    try:
        _dbg_write(
            "together_prompt_preview.txt",
            (
                "=== SYSTEM ===\n"
                "You rewrite WORK EXPERIENCE bullets job-by-job. Keep original job order and EXACT number of jobs.\n"
                'Return ONLY dash bullets for each job ("- ..."). Separate jobs with ONE blank line.\n'
                "Return EXACTLY 5–6 bullets per job. Do NOT include any headers, job titles, companies, dates, or markdown separators.\n\n"
                "=== USER ===\n"
                "Original WORK EXPERIENCE bullets (headers excluded):\n"
                f"{_redact_pii(base_we_for_model).strip()}\n\n"
                "Job Description (focus terms):\n"
                f"{_redact_pii(jd_text).strip()}\n\n"
                f"Priority keywords to weave in: {', '.join(all_keywords[:40])}\n\n"
                "Constraints:\n"
                "- 5–6 bullets per job.\n- Quantify where possible (%, $, #).\n"
                "- No headers, no roles, only bullets.\n- One blank line between jobs.\n"
            )
        )
    except Exception as _e:
        if DEBUG: print(f"⚠️ Could not write together_prompt_preview.txt: {_e}")

    stats = {
        "base_we_raw": _quick_stats("base_we_raw", base_we),
        "base_we_normalized": _quick_stats("base_we_normalized", base_we_normalized),
        "jd_text": _quick_stats("jd_text", jd_text),
        "keywords_count": len(all_keywords),
        "headers_count": len(base_headers),
    }
    _dbg_write("together_preflight_stats.json", json.dumps(stats, indent=2))

    # 2b) TOGETHER CALL — job-aware attempts + cleaning + fallback
    def _looks_empty_or_separators(txt: str) -> bool:
        if not txt or not txt.strip():
            return True
        stripped = re.sub(r"[-—\s*_#>|`~]+", "", txt, flags=re.MULTILINE)
        stripped = re.sub(r"[•\-\*\d\.\(\)\[\]]+", "", stripped)
        return len(stripped.strip()) < 10

    expected_jobs = len(base_headers)
    bullet_model_text = ""
    max_attempts = 2  # soft + hard

    for i in range(1, max_attempts + 1):
        try:
            candidate = _call_bullet_model(together_client, base_we_for_model, jd_text, all_keywords)
        except Exception as e:
            if DEBUG:
                print(f"⚠️ Together call failed (attempt {i}): {e}")
            candidate = ""

        _dbg_write(f"together_bullet_raw_attempt{i}.txt", candidate)

        # Clean instruction junk + markdown/separators
        candidate = _clean_instruction_lines(candidate)
        candidate = _clean_model_block(candidate)
        candidate = _strip_md_separators(candidate)
        candidate = _keep_only_bullets_and_blank(candidate)

        counts = _split_jobs_and_count_bullets(candidate)
        if counts:
            total_bullets = sum(counts)
        else:
            total_bullets = sum(1 for ln in candidate.splitlines() if ln.strip().startswith("- "))

        # Log the fully cleaned attempt BEFORE any early returns
        _dbg_write(f"together_bullet_clean_attempt{i}.txt", candidate)

        # Accept perfect shape (correct job count + 5–6 bullets each)
        if not _looks_empty_or_separators(candidate) and _jobwise_ok(candidate, expected_jobs, lo=5, hi=6):
            bullet_model_text = candidate
            break

        # Accept single-block mega-output; we'll rebalance across jobs later
        if len(counts) == 1 and total_bullets >= expected_jobs * 5:
            _dbg_write(f"together_bullet_accept_reason_attempt{i}.txt", "accepted_single_block_rebalance")
            bullet_model_text = candidate
            break

        _dbg_write(f"together_bullet_clean_attempt{i}.txt", candidate)

        # Validate: non-empty, correct job count, 5–6 bullets per job
        if not _looks_empty_or_separators(candidate) and _jobwise_ok(candidate, expected_jobs, lo=5, hi=6):
            bullet_model_text = candidate
            break

        # One hard retry: bias the model by appending tight constraints into the input text
        if i < max_attempts:
            base_we_for_model = (
                base_we_for_model
                + "\n\n- Return exactly 5–6 bullets per job.\n- No prefaces. No headers.\n"
            )

    _dbg_write("together_bullet_output.txt", bullet_model_text or "")

    # Prefer model candidate: rebuild it first; fallback to base only if rebuild is unusable
    if bullet_model_text:
        rebuilt_we = fix_work_experience_layout(bullet_model_text, base_text)

        # (Optional but recommended) sanity check: ensure 5–6 bullets per job
        if DEBUG and rebuilt_we:
            per_job, cur = [], 0
            for ln in rebuilt_we.splitlines():
                if is_job_header(ln):
                    if cur:
                        per_job.append(cur)
                    cur = 0
                elif ln.strip().startswith("- "):
                    cur += 1
            if cur:
                per_job.append(cur)
            print(f"🧮 Final bullets per job: {per_job} (target: 5–6 each)")

            # If any job is severely underfilled, fall back to base headers
            if any(c < 5 for c in per_job):
                if DEBUG:
                    print("🛟 Rebuilt output too small — falling back to base bullets.")
                fallback_text = "\n".join(base_we.splitlines())
                rebuilt_we = fix_work_experience_layout(fallback_text, base_text)
    else:
        fallback_text = "\n".join(base_we.splitlines())
        rebuilt_we = fix_work_experience_layout(fallback_text, base_text)

    _dbg_write("work_experience_rebuilt.txt", rebuilt_we if rebuilt_we else "")

        # Clean WE bullets locally (do NOT run globally)
    rebuilt_we = remove_placeholder_bullets(rebuilt_we)

# 3) Rebuild doc: replace SUMMARY and WORK EXPERIENCE; preserve everything else
    final_text = base_text
    if summary_text:
        final_text = insert_or_replace_section(final_text, "SUMMARY", summary_text)

    # SAFE whole-section replace for WE
    final_text = _hard_replace_section(final_text, "WORK EXPERIENCE", rebuilt_we)

    # CERTIFICATIONS
    base_certs = _strip_placeholder_block(sections.get("CERTIFICATIONS", "") or "")
    if base_certs:
        final_text = insert_or_replace_section(final_text, "CERTIFICATIONS", base_certs)

    # EDUCATION guard
    if force_education and "EDUCATION" not in extract_resume_sections(final_text):
        final_text = insert_or_replace_section(final_text, "EDUCATION", sections.get("EDUCATION", ""))

    # 5) Header restore & cleanup + spacing hygiene
    # (header repair stays disabled)
    # try:
    #     expected = _collect_base_headers(base_text)
    #     final_text = restore_missing_company_headers(final_text, expected)
    # except Exception:
    #     pass

    # 🔒 Freeze WE while cleaning the rest
    _we_body = rebuilt_we or ""
    final_text = re.sub(
        r"(?ms)^WORK EXPERIENCE\s*\n.*?(\n(?=[A-Z][A-Z &/]+\s*$)|\Z)",
        "WORK EXPERIENCE\n\n<<WE_PLACEHOLDER>>\\1",
        final_text,
    )

    # Run global cleanups outside WE
    final_text = remove_hallucinated_titles(final_text)
    final_text = fix_broken_headers(final_text)
    final_text = _ensure_blank_lines_between_sections(final_text)

    # Restore WE exactly as rebuilt (preserve block as-is)
    final_text = final_text.replace(
        "<<WE_PLACEHOLDER>>",
        _we_body if _we_body.endswith("\n") else _we_body + "\n"
    )

    # Final guard: ensure WE section equals the rebuilt block
    _we_now = _extract_section(final_text, "WORK EXPERIENCE")
    if _norm_block(_we_now) != _norm_block(rebuilt_we):
        if DEBUG:
            print("♻️ WE mismatch after cleanups — hard replacing again.")
        final_text = _hard_replace_section(final_text, "WORK EXPERIENCE", rebuilt_we)

    return final_text



# ---------------- Legacy wrapper (keeps your old callsite) ----------------
from extract_job_keywords import extract_job_info
import requests

def _load_docx_text(path: str) -> str:
    try:
        doc = Document(path); 
        return "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
    except Exception:
        try:
            with open(path, "r", encoding="utf-8") as f: return f.read()
        except Exception: return ""

def _dedupe_keywords(ats_keywords, skills, verbs, jd_text: str) -> List[str]:
    seen, out = set(), []
    bundles = [[kw.strip() for kw in (ats_keywords or []) if kw and kw.strip()], skills or [], verbs or []]
    for bundle in bundles:
        for kw in bundle:
            low = kw.lower()
            if low and low not in seen:
                seen.add(low); out.append(kw)
    for soft in ["communication","teamwork","problem-solving","adaptability","critical thinking","time management","leadership","attention to detail"]:
        if soft in jd_text.lower() and soft not in seen:
            seen.add(soft); out.append(soft)
    return out

class _GroqClientShim:
    def complete(self, model: str, prompt: str) -> str:
        key = os.getenv("GROQ_API_KEY","")
        if not key:
            return ""
        url = "https://api.groq.com/openai/v1/chat/completions"
        payload = {
            "model": os.getenv("GROQ_MODEL","llama-3.1-8b-instant"),
            "messages":[{"role":"user","content":prompt}],
            "temperature":0.3,
        }
        headers = {"Authorization": f"Bearer {key}"}
        # retries on transient 5xx/429 with exponential backoff
        for attempt in range(1, 4):
            try:
                r = requests.post(url, headers=headers, json=payload, timeout=60)
                status = r.status_code
                if status >= 500 or status == 429:
                    raise RuntimeError(f"Groq transient status {status}")
                r.raise_for_status()
                j = r.json()
                out = (j.get("choices") or [{}])[0].get("message", {}).get("content", "") or ""
                if out.strip():
                    return out
                # empty content: treat as soft failure
                raise RuntimeError("Groq returned empty content")
            except Exception as e:
                try:
                    _dbg_write(f"groq_summary_attempt{attempt}.txt", f"{type(e).__name__}: {e}")
                except Exception:
                    pass
                if attempt < 3:
                    time.sleep(1.5 ** attempt)
                    continue
                return ""  # swallow and let callers fall back

class _TogetherClientShim:
    def chat(self, system: str, user: str, model: str) -> str:
        try:
            from models import call_bullet_model
        except Exception:
            call_bullet_model = None
        prompt = f"{system}\n\n{user}"
        if call_bullet_model:
            return call_bullet_model(prompt, max_tokens=3200)
        return ""

def tailor_resume_with_models(
    job_title: str,
    job_description: str,
    *,
    base_resume_path: str = "data/base_resume_test.txt",
    ats_keywords: List[str] | None = None,
    ats_sections: List[str] | None = None,
) -> str:
    base_text = load_resume_text(base_resume_path)

    # SAFE: extract_job_info may fail; tolerate gracefully
    job_info = {}
    try:
        maybe = extract_job_info(f"Job Title: {job_title}\n\n{job_description}")
        if isinstance(maybe, dict):
            job_info = maybe
    except Exception as e:
        _dbg_write("extract_job_info_error.txt", f"{type(e).__name__}: {e}")

    # Build + sanitize keywords
    all_keywords = _dedupe_keywords(ats_keywords or [], job_info.get("skills", []), job_info.get("verbs", []), job_description)
    all_keywords, dropped = _sanitize_keywords(all_keywords)
    _dbg_write("keywords_used.txt", "\n".join(all_keywords))
    if dropped:
        _dbg_write("keywords_dropped.txt", "\n".join(dropped))

    groq_client, together_client = _GroqClientShim(), _TogetherClientShim()
    return _tailor_resume_with_models_core(
        base_text=base_text,
        jd_text=job_description,
        all_keywords=all_keywords,
        groq_client=groq_client,
        together_client=together_client,
        force_education=True,
    )

# ---------------------- Export for bulk script ----------------------
def save_to_docx(text: str, path: str):
    from docx.shared import Pt
    doc = Document()
    VALID = {"SUMMARY","SKILLS","WORK EXPERIENCE","CERTIFICATIONS","EDUCATION"}
    for line in text.splitlines():
        clean = line.strip().replace("•","-").replace("·","-")
        clean = clean.encode("utf-8","ignore").decode("utf-8").replace("\n"," ").strip()
        upper = re.sub(r"[^A-Z]", "", clean.upper())
        if upper in {h.replace(" ","") for h in VALID}:
            para = doc.add_paragraph(); run = para.add_run(next(h for h in VALID if upper == h.replace(" ","")))
            run.bold = True; para.paragraph_format.space_after = Pt(8)
            continue
        if clean:
            para = doc.add_paragraph(clean); para.paragraph_format.space_after = Pt(6)
        else:
            doc.add_paragraph("")
    doc.save(path)

# Optional convenience
def parse_sections(text: str) -> Dict[str,str]:
    return extract_resume_sections(text)